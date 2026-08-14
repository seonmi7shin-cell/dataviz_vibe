# output/그래프_분석_보고서.md 의 내용을 그대로 워드(docx) 문서로 만드는 스크립트.
# 이미 저장된 output/charts 의 그래프 이미지를 그대로 첨부하고, 표·분석 문구를 docx 표/문단으로 옮긴다.

import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.chdir(os.path.dirname(os.path.abspath(__file__)) + "/..")

CHART_DIR = os.path.join("output", "charts")

# (제목, 이미지 파일명, 표 헤더, 표 데이터, 시사점) 순서로 각 섹션을 정의.
# 표 데이터·시사점은 output/그래프_분석_보고서.md 와 동일한 내용을 그대로 옮긴 것이다.
sections = [
    {
        "title": "1. 가맹점업종별 거래금액 합계",
        "image": "업종별_거래금액_가로막대.png",
        "headers": ["업종", "합계(원)"],
        "rows": [
            ["여행", "127,236,638"], ["쇼핑", "122,086,016"], ["교육", "104,255,218"],
            ["공과금", "71,544,296"], ["식음료", "48,439,587"], ["문화/여가", "41,583,584"],
            ["의료", "40,060,538"], ["기타", "21,167,883"], ["교통", "5,307,392"],
        ],
        "insight": (
            "여행·쇼핑·교육 상위 3개 업종이 전체 거래금액의 절반 이상을 차지한다. "
            "교통은 결제 건수는 많지만(1,485건) 건당 금액이 작아 총액 기준으로는 가장 낮다 — "
            "즉 총액 순위는 결제 빈도가 아니라 건당 단가에 좌우된다."
        ),
    },
    {
        "title": "2. 지역별 거래금액 합계",
        "image": "지역별_거래금액_가로막대.png",
        "headers": None,
        "rows": None,
        "insight": (
            "서울(163.2M)과 경기(139.2M) 두 지역이 전체의 절반 이상을 차지하는 수도권 쏠림 현상이 "
            "뚜렷하다. 세종(5.5M)이 가장 낮은데, 이는 인구·상권 규모 자체가 작기 때문으로 보인다."
        ),
    },
    {
        "title": "3. 결제수단별 거래금액 합계",
        "image": "결제수단별_거래금액_가로막대.png",
        "headers": None,
        "rows": None,
        "insight": (
            "신용카드가 압도적 1위다. 다만 이 그래프는 원본 표기(공백 포함 ' 신용카드' 별도 항목) 그대로 "
            "집계한 것이라 신용카드 실제 비중이 과소평가되어 있다 — 정확한 비중은 4·5번 도넛/비율 차트를 참고."
        ),
    },
    {
        "title": "4. 결제수단별 건수 및 비율",
        "image": "결제수단별_건수_비율.png",
        "headers": ["결제수단", "건수", "비율"],
        "rows": [
            ["신용카드", "5,385", "44.32%"], ["체크카드", "2,599", "21.39%"],
            ["계좌이체", "1,930", "15.88%"], ["포인트결제", "1,130", "9.30%"],
            ["휴대폰결제", "761", "6.26%"], ["결측", "345", "2.84%"],
        ],
        "insight": (
            "공백 표기 문제를 정리(strip())하고 나면 신용카드가 건수 기준 전체의 44%로 나머지 네 수단을 "
            "합친 것과 비슷한 규모다. 카드 기반 결제(신용+체크)가 전체의 66%에 달한다."
        ),
    },
    {
        "title": "5. 결제수단별 거래금액 비중 (도넛차트)",
        "image": "결제수단별_거래금액_도넛차트.png",
        "headers": ["결제수단", "합계(원)", "비중", "건수"],
        "rows": [
            ["신용카드", "254,500,813", "43.75%", "5,385"],
            ["체크카드", "124,975,625", "21.49%", "2,599"],
            ["계좌이체", "93,718,496", "16.11%", "1,930"],
            ["포인트결제", "54,344,379", "9.34%", "1,130"],
            ["휴대폰결제", "35,685,922", "6.13%", "761"],
            ["결측", "18,455,917", "3.17%", "345"],
            ["전체", "581,681,927", "100%", "12,150"],
        ],
        "insight": (
            "도넛 조각마다 비율과 건수를 함께 표시했고, 가운데에는 전체 건수(12,150건)를 넣었다. "
            "건수 비중(44.32%)과 금액 비중(43.75%)이 거의 같다 — 신용카드는 건당 평균 금액도 다른 "
            "수단과 비슷한 수준이라는 뜻이다. 결측 거래도 1,845만원 규모라 무시할 수 없는 금액이므로, "
            "결측 처리 방식이 총액 집계에 영향을 준다."
        ),
    },
    {
        "title": "6. 연령대별 거래금액 합계",
        "image": "연령대별_거래금액_가로막대.png",
        "headers": ["연령대", "합계(원)"],
        "rows": [
            ["30대", "158,357,914"], ["20대", "150,427,675"], ["40대", "104,277,043"],
            ["50대", "70,323,782"], ["60대 이상", "33,895,591"], ["10대", "27,137,422"],
        ],
        "insight": (
            "20~30대가 전체 거래금액의 절반 가까이를 차지해 핵심 소비층임을 보여준다. "
            "연령대가 올라갈수록(40대 이후) 완만하게 감소하는 패턴이다."
        ),
    },
    {
        "title": "7. 연월별 거래금액 합계 추이",
        "image": "연월별_거래금액_라인차트.png",
        "headers": None,
        "rows": None,
        "insight": (
            "월별 합계가 43.5M~55.2M원 사이에서 등락하며 뚜렷한 상승/하락 추세나 계절성은 보이지 "
            "않는다. 최댓값은 3월(55.2M), 최솟값은 10월(43.5M)로 큰 차이는 아니라 월별 변동은 안정적인 편이다."
        ),
    },
]

doc = Document()
doc.styles["Normal"].font.size = Pt(10)

section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

title = doc.add_heading("핀테크 결제 데이터 그래프 분석 보고서", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("데이터: data/01_핀테크결제_dirty.csv (원본, 12,150건)")

for sec in sections:
    doc.add_heading(sec["title"], level=2)
    doc.add_picture(os.path.join(CHART_DIR, sec["image"]), width=Cm(15))

    if sec["headers"]:
        table = doc.add_table(rows=1, cols=len(sec["headers"]))
        table.style = "Light Grid Accent 1"
        for i, name in enumerate(sec["headers"]):
            table.rows[0].cells[i].text = name
        for row_values in sec["rows"]:
            cells = table.add_row().cells
            for i, value in enumerate(row_values):
                cells[i].text = value

    p = doc.add_paragraph()
    run = p.add_run("시사점: ")
    run.bold = True
    p.add_run(sec["insight"])

doc.add_heading("종합 요약", level=2)
doc.add_paragraph(
    "업종·지역·연령대 모두 뚜렷한 쏠림(여행/쇼핑, 서울/경기, 20~30대)이 있어 상위 카테고리 중심 "
    "전략이 유효해 보인다. 결제수단은 신용카드 비중이 압도적이지만, 원본 데이터의 공백 표기 문제 "
    "때문에 정리 전/후 집계 결과가 달라진다는 점을 유의해야 한다. 시간 흐름(연월)에는 특별한 추세가 "
    "없어, 계절성보다는 카테고리·지역·연령 요인이 거래금액을 설명하는 더 중요한 변수로 보인다."
)

output_path = os.path.join("output", "그래프_분석_보고서.docx")
doc.save(output_path)
print(f"저장 경로: {output_path}")
