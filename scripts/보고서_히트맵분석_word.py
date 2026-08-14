# output/지역업종_히트맵_분석보고서.md 의 내용을 그대로 워드(docx) 문서로 만드는 스크립트.

import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.chdir(os.path.dirname(os.path.abspath(__file__)) + "/..")

doc = Document()
doc.styles["Normal"].font.size = Pt(10)

section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

title = doc.add_heading("지역×업종 거래금액 히트맵 분석 보고서", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("데이터: data/핀테크_정제완료.csv (정제 완료, 11,738건)")

doc.add_picture(os.path.join("output", "charts", "지역업종_거래금액_히트맵.png"), width=Cm(16))

doc.add_heading("1. 전체 구조", level=2)
doc.add_paragraph(
    "지역 17개(실제 지역, 가나다순) × 업종 9개를 교차 집계했다. '알수없음'(지역 결측 대체값, "
    "475건 유래)은 실제 지역과 순위 비교에 섞이지 않도록 히트맵 맨 아래 구분선 밑에 별도로 배치했다."
)

doc.add_heading("2. 가장 진한 칸 / 가장 옅은 칸", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
for i, name in enumerate(["구분", "조합", "거래금액 합계"]):
    table.rows[0].cells[i].text = name
for row in [
    ["최댓값 (가장 진한 칸)", "서울 × 여행", "39,116,569원"],
    ["최솟값 (0 제외, 가장 옅은 칸)", "세종 × 교통", "51,434원"],
]:
    cells = table.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v
doc.add_paragraph("두 값의 차이가 760배에 달할 만큼 격차가 크다 — 지역과 업종 조합에 따라 거래 규모가 극단적으로 갈린다.")

doc.add_heading("3. 지역별 총합 (상위 3 · 하위 1)", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
for i, name in enumerate(["순위", "지역", "합계(원)"]):
    table.rows[0].cells[i].text = name
for row in [
    ["1", "서울", "160,603,512"],
    ["2", "경기", "137,057,485"],
    ["3", "부산", "37,150,102"],
    ["17", "세종", "5,510,049"],
]:
    cells = table.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v
doc.add_paragraph("서울·경기 두 지역이 나머지 15개 지역을 합친 것과 맞먹는 규모다 (전체의 약 55%).")

doc.add_heading("4. 업종별 총합 (상위 3 · 하위 1)", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
for i, name in enumerate(["순위", "업종", "합계(원)"]):
    table.rows[0].cells[i].text = name
for row in [
    ["1", "여행", "120,089,719"],
    ["2", "쇼핑", "117,001,680"],
    ["3", "교육", "98,439,821"],
    ["9", "교통", "5,013,118"],
]:
    cells = table.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v

doc.add_heading("5. '알수없음'(지역 결측) 행 별도 확인", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
for i, name in enumerate(["업종", "합계(원)"]):
    table.rows[0].cells[i].text = name
for row in [
    ["여행", "6,264,289"], ["교육", "3,900,570"], ["쇼핑", "3,597,677"],
    ["공과금", "2,796,563"], ["문화/여가", "2,041,567"], ["식음료", "1,811,551"],
    ["의료", "1,301,924"], ["기타", "683,754"], ["교통", "231,158"],
]:
    cells = table.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v
doc.add_paragraph(
    "지역을 알 수 없는 거래도 업종별 분포는 실제 지역들과 비슷한 패턴(여행·교육·쇼핑이 상위)을 "
    "보인다 — 즉 결측이 특정 업종에 쏠려서 생긴 게 아니라 무작위에 가깝게 발생한 것으로 보인다."
)

doc.add_heading("종합 시사점", level=2)
points = [
    "수도권(서울·경기) 집중이 매우 강하다. 두 지역만으로 전체 거래금액의 절반을 넘는다. "
    "지역별 마케팅·리소스 배분을 고려한다면 수도권 외 지역은 상대적으로 작은 시장으로 봐야 한다.",
    "여행·쇼핑 업종이 어느 지역에서나 상위권이지만, 특히 서울에서 여행 업종 거래금액(39.1M)이 "
    "다른 어떤 지역×업종 조합보다도 압도적으로 크다 — 서울 지역 여행 관련 결제가 이 데이터셋의 '핫스팟'이다.",
    "교통 업종은 모든 지역에서 가장 낮은 수준이라 지역차보다 업종 자체의 특성(건당 소액 결제)이 "
    "지배적인 요인으로 보인다.",
    "결측(알수없음) 거래의 업종 분포가 정상 지역들과 유사해, 이 결측이 특정 상황(예: 특정 업종 전산 "
    "오류)이 아니라 전반적으로 고르게 발생했을 가능성을 시사한다.",
]
for p in points:
    doc.add_paragraph(p, style="List Number")

output_path = os.path.join("output", "지역업종_히트맵_분석보고서.docx")
doc.save(output_path)
print(f"저장 경로: {output_path}")
