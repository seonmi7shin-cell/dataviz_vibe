# 가맹점업종별 거래금액 분석을 한 페이지짜리 docx 보고서로 만드는 스크립트.
# 표·차트를 넣는 반복 작업이라 python-docx 로 문서를 조립하고, 차트는 이미 저장해둔
# output/charts/가맹점업종_거래금액.png 를 그대로 첨부한다.

import os
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.chdir(os.path.dirname(os.path.abspath(__file__)) + "/..")

df = pd.read_csv(os.path.join("data", "핀테크_정제완료.csv"), encoding="utf-8-sig")
chart_path = os.path.join("output", "charts", "가맹점업종_거래금액.png")

by_cat = df.groupby("가맹점업종")["거래금액"].agg(합계="sum", 평균="mean", 건수="count")
by_cat = by_cat.sort_values("합계", ascending=False)
total = by_cat["합계"].sum()
by_cat["비중(%)"] = (by_cat["합계"] / total * 100).round(2)

top_name = by_cat.index[0]
top_row = by_cat.iloc[0]
bottom_name = by_cat.index[-1]
bottom_row = by_cat.iloc[-1]

doc = Document()

# 한 페이지에 맞추려고 여백을 좁게 잡음
section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

title = doc.add_heading("가맹점업종별 거래금액 분석 보고서", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(chart_path, width=Cm(16))
doc.styles["Normal"].font.size = Pt(10)

doc.add_heading("분석 요약", level=2)
summary = doc.add_paragraph()
summary.add_run(
    f"정제 완료 데이터(총 {len(df):,}건) 기준으로 가맹점업종별 거래금액 합계를 집계한 결과, "
    f"'{top_name}' 업종이 {top_row['합계']:,.0f}원으로 전체({total:,.0f}원)의 "
    f"{top_row['비중(%)']:.1f}%를 차지해 가장 큰 비중을 보였다. "
    f"건당 평균 결제금액은 {top_row['평균']:,.0f}원으로, 건수({top_row['건수']:,}건) 자체는 "
    f"'쇼핑'이나 '식음료'보다 적지만 건당 단가가 높아 총액이 1위로 올라온 것으로 보인다. "
    f"반대로 '{bottom_name}' 업종은 {bottom_row['합계']:,.0f}원으로 전체의 "
    f"{bottom_row['비중(%)']:.1f}%에 그쳐, 결제 건수({bottom_row['건수']:,}건)는 많지만 "
    f"건당 금액이 작아 총액 기준으로는 가장 낮게 나타났다."
)

doc.add_heading("업종별 상세 지표", level=2)
table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, name in enumerate(["가맹점업종", "합계(원)", "평균(원)", "건수", "비중(%)"]):
    hdr[i].text = name

for name, row in by_cat.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(name)
    cells[1].text = f"{row['합계']:,.0f}"
    cells[2].text = f"{row['평균']:,.0f}"
    cells[3].text = f"{int(row['건수']):,}"
    cells[4].text = f"{row['비중(%)']:.2f}"

output_path = os.path.join("output", "가맹점업종_분석보고서.docx")
os.makedirs("output", exist_ok=True)
doc.save(output_path)
print(f"보고서 저장 경로: {output_path}")
print(f"1위 업종: {top_name} ({top_row['비중(%)']:.1f}%)")
print(f"최하위 업종: {bottom_name} ({bottom_row['비중(%)']:.1f}%)")
